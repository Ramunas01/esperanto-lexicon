"""Tests for src/ingestion/docx_to_corpus.py.

Documents are constructed programmatically with python-docx so tests require
no external corpus files and work in a checked-out-code-only environment.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

import sys
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from ingestion.docx_to_corpus import (
    convert,
    extract_definition,
    is_amendment,
    is_heading,
    is_table_note,
)


# ---------------------------------------------------------------------------
# Helpers for building synthetic paragraphs / documents
# ---------------------------------------------------------------------------


def _make_doc() -> Document:
    return Document()


def _add_plain(doc: Document, text: str):
    return doc.add_paragraph(text)


def _add_italic(doc: Document, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = True
    return para


def _add_bold_heading(doc: Document, text: str):
    """All-bold short paragraph — classified as HEADING."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    return para


def _add_definition(doc: Document, term: str, definition: str):
    """Bold term run followed by ' – definition' run."""
    para = doc.add_paragraph()
    bold_run = para.add_run(term)
    bold_run.bold = True
    para.add_run(f" – {definition}")
    return para


def _add_style_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _save_and_load(doc: Document, tmp_path: Path) -> Path:
    path = tmp_path / "test_input.docx"
    doc.save(str(path))
    return path


def _run_convert(doc: Document, tmp_path: Path) -> tuple[str, str]:
    """Save doc, convert, return (main_text, amendments_text)."""
    input_path = _save_and_load(doc, tmp_path)
    output_path = tmp_path / "output.txt"
    amendments_path = tmp_path / "amendments.txt"
    convert(input_path, output_path, amendments_path)
    main_text = output_path.read_text(encoding="utf-8")
    amendments_text = amendments_path.read_text(encoding="utf-8")
    return main_text, amendments_text


# ---------------------------------------------------------------------------
# Unit tests: classifiers
# ---------------------------------------------------------------------------


class TestIsAmendment:
    def test_italic_only_is_amendment(self) -> None:
        doc = _make_doc()
        para = _add_italic(doc, "Nr. XIII-841, 2017-12-07, paskelbta TAR")
        assert is_amendment(para) is True

    def test_nr_reference_is_amendment(self) -> None:
        doc = _make_doc()
        para = _add_plain(doc, "Nr. XIII-841, 2017-12-07, paskelbta TAR")
        assert is_amendment(para) is True

    def test_paskelbta_tar_is_amendment(self) -> None:
        doc = _make_doc()
        para = _add_plain(doc, "2017-12-07, paskelbta TAR je 2017-12-14")
        assert is_amendment(para) is True

    def test_straipsnio_pakeitimai_is_amendment(self) -> None:
        doc = _make_doc()
        para = _add_plain(doc, "Straipsnio pakeitimai:")
        assert is_amendment(para) is True

    def test_straipsnio_dalies_pakeitimai_is_amendment(self) -> None:
        doc = _make_doc()
        para = _add_plain(doc, "Straipsnio dalies pakeitimai:")
        assert is_amendment(para) is True

    def test_straipsnio_punkto_pakeitimai_is_amendment(self) -> None:
        doc = _make_doc()
        para = _add_plain(doc, "Straipsnio punkto pakeitimai:")
        assert is_amendment(para) is True

    def test_normal_paragraph_not_amendment(self) -> None:
        doc = _make_doc()
        para = _add_plain(doc, "Gyventojas – asmuo, nuolat gyvenantis Lietuvoje.")
        assert is_amendment(para) is False


class TestIsTableNote:
    def test_tar_pastaba_is_table_note(self) -> None:
        assert is_table_note("TAR pastaba. Žr. 2017 m. gruodžio 14 d.") is True

    def test_normal_text_not_table_note(self) -> None:
        assert is_table_note("Gyventojas") is False

    def test_tar_pastaba_must_be_at_start(self) -> None:
        assert is_table_note("Pastaba: TAR pastaba yra čia") is False


class TestIsHeading:
    def test_style_heading_detected(self) -> None:
        doc = _make_doc()
        _add_style_heading(doc, "7 straipsnis. Gyventojai", level=1)
        # add_heading creates a paragraph after the default blank one
        para = next(p for p in doc.paragraphs if p.text.strip())
        assert is_heading(para) is True

    def test_all_bold_short_no_dash_is_heading(self) -> None:
        doc = _make_doc()
        para = _add_bold_heading(doc, "7 straipsnis")
        assert is_heading(para) is True

    def test_all_bold_with_dash_not_heading(self) -> None:
        doc = _make_doc()
        para = _add_definition(doc, "Gyventojas", "nuolat gyvenantis asmuo")
        assert is_heading(para) is False

    def test_long_bold_text_not_heading(self) -> None:
        doc = _make_doc()
        para = doc.add_paragraph()
        run = para.add_run("A" * 90)
        run.bold = True
        assert is_heading(para) is False


class TestExtractDefinition:
    def test_bold_term_wrapped_in_markers(self) -> None:
        doc = _make_doc()
        para = _add_definition(doc, "Gyventojas", "nuolat gyvenantis asmuo")
        result = extract_definition(para)
        assert result is not None
        assert result.startswith("**Gyventojas**")
        assert " – nuolat gyvenantis asmuo" in result

    def test_no_dash_returns_none(self) -> None:
        doc = _make_doc()
        para = _add_plain(doc, "Paprastas tekstas be brūkšnelio.")
        assert extract_definition(para) is None

    def test_plain_hyphen_not_a_definition_separator(self) -> None:
        # Mirrors English heading "13 (1) - Determination and taxation..."
        # Plain hyphen (-) must not trigger definition detection.
        doc = _make_doc()
        para = doc.add_paragraph()
        run = para.add_run("13 (1)")
        run.bold = True
        para.add_run(" - Determination and taxation of income.")
        assert extract_definition(para) is None

    def test_dash_but_no_bold_returns_none(self) -> None:
        doc = _make_doc()
        para = _add_plain(doc, "normalus tekstas – be bold")
        assert extract_definition(para) is None

    def test_bold_with_clause_number(self) -> None:
        doc = _make_doc()
        para = doc.add_paragraph()
        run = para.add_run("1. Darbo santykiai")
        run.bold = True
        para.add_run(" – santykiai tarp darbdavio ir darbuotojo.")
        result = extract_definition(para)
        assert result is not None
        assert "**1. Darbo santykiai**" in result
        assert "santykiai tarp darbdavio" in result

    def test_bold_term_over_120_chars_not_definition(self) -> None:
        doc = _make_doc()
        para = doc.add_paragraph()
        # Mirrors Art. 3 §3: "3. Mokesčio administratorius nuolatinio Lietuvos
        # gyventojo prašymu, mokestiniam laikotarpiui pasibaigus, Vyriausybės
        # arba jos įgaliotos institucijos nustatyta tvarka privalo pervesti..."
        # (~200 chars of bold before the dash)
        long_bold = (
            "3. Mokesčio administratorius nuolatinio Lietuvos gyventojo "
            "prašymu, mokestiniam laikotarpiui pasibaigus, Vyriausybės "
            "arba jos įgaliotos institucijos nustatyta tvarka"
        )
        run = para.add_run(long_bold)
        run.bold = True
        para.add_run(" – privalo pervesti asmenims.")
        assert len(long_bold) > 120, f"test string must be >120 chars, got {len(long_bold)}"
        assert extract_definition(para) is None

    def test_bold_term_within_120_chars_is_definition(self) -> None:
        doc = _make_doc()
        para = doc.add_paragraph()
        run = para.add_run("Kontroliuojamasis vienetas")
        run.bold = True
        para.add_run(" – užsienio juridinis asmuo.")
        result = extract_definition(para)
        assert result is not None
        assert "**Kontroliuojamasis vienetas**" in result

    def test_bold_term_with_comma_is_still_definition(self) -> None:
        # Comma in the term does NOT disqualify — only length matters now.
        doc = _make_doc()
        para = doc.add_paragraph()
        run = para.add_run("1. Rezidentas, užsienietis")
        run.bold = True
        para.add_run(" – asmuo gyvenantis Lietuvoje.")
        result = extract_definition(para)
        assert result is not None
        assert "**1. Rezidentas, užsienietis**" in result


# ---------------------------------------------------------------------------
# Integration tests: convert()
# ---------------------------------------------------------------------------


class TestAmendmentsExcludedFromMain:
    def test_italic_amendment_not_in_output(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_plain(doc, "Normalus tekstas.")
        _add_italic(doc, "Nr. XIII-841, 2017-12-07, paskelbta TAR je 2017-12-14")
        main, _ = _run_convert(doc, tmp_path)
        assert "Nr. XIII-841" not in main

    def test_nr_reference_not_in_output(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_plain(doc, "Tikras turinys.")
        _add_plain(doc, "Nr. XIV-1697, 2022-12-15, paskelbta TAR")
        main, _ = _run_convert(doc, tmp_path)
        assert "Nr. XIV-1697" not in main

    def test_amendment_written_to_amendments_file(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_plain(doc, "Tikras turinys.")
        _add_italic(doc, "Nr. XIII-841, 2017-12-07, paskelbta TAR je 2017-12-14")
        _, amendments = _run_convert(doc, tmp_path)
        assert "Nr. XIII-841" in amendments
        assert "2017-12-07" in amendments

    def test_amendments_file_format(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_italic(doc, "Nr. XIII-841, 2017-12-07, paskelbta TAR")
        _, amendments = _run_convert(doc, tmp_path)
        lines = [ln for ln in amendments.splitlines() if ln.strip()]
        assert len(lines) == 1
        parts = lines[0].split(" | ")
        assert len(parts) == 3  # ARTICLE | REF | DATE

    def test_amendment_date_extracted(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_italic(doc, "Nr. XIII-841, 2017-12-07, paskelbta TAR")
        _, amendments = _run_convert(doc, tmp_path)
        assert "2017-12-07" in amendments

    def test_normal_content_stays_in_output(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_plain(doc, "Gyventojas yra asmuo.")
        _add_italic(doc, "Nr. XIII-841, 2017-12-07, paskelbta TAR")
        main, _ = _run_convert(doc, tmp_path)
        assert "Gyventojas yra asmuo." in main

    def test_straipsnio_pakeitimai_excluded(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_plain(doc, "Turinys.")
        _add_plain(doc, "Straipsnio pakeitimai:")
        main, _ = _run_convert(doc, tmp_path)
        assert "Straipsnio pakeitimai" not in main

    def test_tar_pastaba_excluded_from_both_files(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_plain(doc, "Turinys.")
        _add_plain(doc, "TAR pastaba. Žr. 2017 m. gruodžio 14 d.")
        main, amendments = _run_convert(doc, tmp_path)
        assert "TAR pastaba" not in main
        assert "TAR pastaba" not in amendments


class TestBoldMarkersPreserved:
    def test_bold_term_has_double_star_markers(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_definition(doc, "Gyventojas", "nuolat gyvenantis asmuo")
        main, _ = _run_convert(doc, tmp_path)
        assert "**Gyventojas**" in main

    def test_definition_separator_preserved(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_definition(doc, "Gyventojas", "nuolat gyvenantis asmuo")
        main, _ = _run_convert(doc, tmp_path)
        assert " – " in main

    def test_definition_text_preserved(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_definition(doc, "Gyventojas", "nuolat gyvenantis asmuo Lietuvoje")
        main, _ = _run_convert(doc, tmp_path)
        assert "nuolat gyvenantis asmuo Lietuvoje" in main

    def test_plain_paragraph_has_no_star_markers(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_plain(doc, "Normalus tekstas be apibrėžimo.")
        main, _ = _run_convert(doc, tmp_path)
        assert "**" not in main

    def test_multiple_definitions_all_marked(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_definition(doc, "Gyventojas", "asmuo gyvenantis Lietuvoje")
        _add_definition(doc, "Pajamos", "gautos lėšos")
        main, _ = _run_convert(doc, tmp_path)
        assert "**Gyventojas**" in main
        assert "**Pajamos**" in main


class TestEmptyParagraphsSkipped:
    def test_empty_paragraph_not_in_output(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_plain(doc, "Pirmas.")
        doc.add_paragraph("")  # explicit empty
        _add_plain(doc, "Antras.")
        main, _ = _run_convert(doc, tmp_path)
        lines = [ln for ln in main.splitlines() if ln]
        # Only content lines; blank lines from heading formatting are allowed
        content_lines = [ln for ln in lines if ln.strip()]
        assert len(content_lines) == 2

    def test_whitespace_only_paragraph_skipped(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_plain(doc, "Pirmas.")
        doc.add_paragraph("   ")
        _add_plain(doc, "Antras.")
        main, _ = _run_convert(doc, tmp_path)
        content_lines = [ln for ln in main.splitlines() if ln.strip()]
        assert len(content_lines) == 2


class TestOutputEncoding:
    def test_output_is_valid_utf8(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_plain(doc, "Tekstas su lietuviškais simboliais: ąčęėįšųūž.")
        _add_definition(doc, "Gyventojas", "nuolat gyvenantis asmuo")
        input_path = _save_and_load(doc, tmp_path)
        output_path = tmp_path / "output.txt"
        amendments_path = tmp_path / "amendments.txt"
        convert(input_path, output_path, amendments_path)
        # Read as bytes and decode — must not raise
        raw = output_path.read_bytes()
        decoded = raw.decode("utf-8")
        assert "ąčęėįšųūž" in decoded

    def test_lithuanian_characters_in_definitions(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_definition(doc, "Įmonė", "juridinis asmuo, vykdantis ūkinę veiklą")
        main, _ = _run_convert(doc, tmp_path)
        assert "**Įmonė**" in main
        assert "juridinis asmuo" in main

    def test_esperanto_characters_preserved(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_plain(doc, "Esperanto simboliai: ĉ, ĝ, ĥ, ĵ, ŝ, ŭ.")
        main, _ = _run_convert(doc, tmp_path)
        assert "ĉ, ĝ, ĥ, ĵ, ŝ, ŭ" in main


class TestStats:
    def test_stats_count_definitions(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_definition(doc, "Gyventojas", "asmuo")
        _add_definition(doc, "Pajamos", "lėšos")
        input_path = _save_and_load(doc, tmp_path)
        stats = convert(input_path, tmp_path / "out.txt", tmp_path / "am.txt")
        assert stats["n_definitions"] == 2

    def test_stats_count_amendments(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_italic(doc, "Nr. XIII-841, 2017-12-07, paskelbta TAR")
        _add_italic(doc, "Nr. XIV-100, 2020-01-01, paskelbta TAR")
        input_path = _save_and_load(doc, tmp_path)
        stats = convert(input_path, tmp_path / "out.txt", tmp_path / "am.txt")
        assert stats["n_amendments"] == 2

    def test_stats_count_table_notes(self, tmp_path: Path) -> None:
        doc = _make_doc()
        _add_plain(doc, "TAR pastaba. Žr. 2017 m.")
        input_path = _save_and_load(doc, tmp_path)
        stats = convert(input_path, tmp_path / "out.txt", tmp_path / "am.txt")
        assert stats["n_table_notes"] == 1
